from docx import Document
from docx.shared import Pt
from datetime import date
import os

# Report metadata
start_date = "2026-04-22"
end_date = "2026-05-06"
report_date = date.today().isoformat()

pointers = [
    ("Gateway Integration & Provider Architecture",
     "Designed and validated a multi-provider AI chain (Bedrock → Gateway → local fallback). Ensured reliable routing and failover behavior for desktop builds."),
    ("Desktop Packaging & EXE Stability",
     "Troubleshot PyInstaller and electron-builder; ensured frontend and backend are bundled correctly and data/config folders are created in %APPDATA% on Windows."),
    ("Screen Flicker Root Cause & Fix",
     "Resolved an infinite re-render loop in the NanoFACS AI panel by adding a one-shot auto-run guard to stop repeated retries when AI calls fail."),
    ("Chat Endpoint Rendering Error Fix",
     "Fixed a payload-normalization bug causing TypeError in the Research Chat UI by properly normalizing incoming message parts."),
    ("Environment Variable Robustness",
     "Improved loading of gateway config in `run_desktop.py` so `CRMIT_AI_GATEWAY_URL` and license key are reliably set before app initialization."),
    ("Graceful Fallbacks for AI Endpoints",
     "Implemented structured fallback responses for chat, NTA, and NanoFACS endpoints to avoid 503s and provide helpful local-analysis messages when gateway fails."),
    ("Logging & Diagnostics Enhancements",
     "Added startup and runtime logs (e.g., '[BioVaram] Gateway config loaded from:') to simplify remote troubleshooting on client machines."),
    ("Expanded Metadata Comparison",
     "Enhanced NTA metadata comparison from 3 to 12 fields with tolerances and severity scoring, improving cross-sample comparability checks."),
    ("Build Validation & Automation",
     "Automated build pipeline runs and validations; verified PyInstaller output, frontend build, and packaged installer artifacts to accelerate release cycles."),
    ("Iterative Release & QA Process",
     "Ran iterative releases (0.1.5 → 0.1.6), analyzed tester feedback, and shipped targeted patches—improving reliability for client deployments.")
]

# Create document
doc = Document()

# Title
title = doc.add_heading(f"AI Productivity Report — Past 2 Weeks", level=1)

# Metadata
doc.add_paragraph(f"To: Parvesh")
doc.add_paragraph(f"From: BioVaram Engineering")
doc.add_paragraph(f"Date: {report_date}")
doc.add_paragraph(f"Reporting period: {start_date} to {end_date}")

doc.add_paragraph("")

intro = doc.add_paragraph()
intro.add_run("Summary:").bold = True
intro.add_run(" Over the past two weeks we focused on improving AI reliability and desktop stability for the BioVaram EV Analysis Platform. Below are 10 concise pointers describing what we used AI for, what we fixed, and how we tested changes.")

doc.add_paragraph("")

for idx, (title, detail) in enumerate(pointers, start=1):
    p = doc.add_paragraph()
    p.add_run(f"{idx}. {title}").bold = True
    p.add_run(f" — {detail}")

# Footer / next steps
doc.add_paragraph("")
next_steps = doc.add_paragraph()
next_steps.add_run("Next Steps:").bold = True
next_steps.add_run(" Manual QA on client machine, collect biovaram.log if issues persist, and finalize v0.1.6 GitHub release assets.")

# Ensure output directories
downloads_path = os.path.expanduser(r"C:\Users\Asus\Downloads")
if not os.path.exists(downloads_path):
    downloads_path = os.path.expanduser("~")

out_dir_repo = os.path.join(os.getcwd(), "docs")
os.makedirs(out_dir_repo, exist_ok=True)

file_name = f"AI_PRODUCTIVITY_REPORT_FOR_PARVESH_{end_date}.docx"
repo_path = os.path.join(out_dir_repo, file_name)
user_path = os.path.join(downloads_path, file_name)

# Save both
try:
    doc.save(repo_path)
    doc.save(user_path)
    print(f"Saved report to: {repo_path}")
    print(f"Saved report to: {user_path}")
except Exception as e:
    print(f"Failed to save report: {e}")
    raise
